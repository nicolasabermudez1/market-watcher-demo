"""Document generation tools — Word digest and cert-chase email."""

import json
from datetime import datetime
from pathlib import Path

from agents import function_tool
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from market_watcher.config import (
    CENTRICA_LAVENDER, CENTRICA_MINT, CENTRICA_NAVY,
    CENTRICA_PALE_LAV, CENTRICA_PURPLE, OUTPUTS_DIR,
)


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _set_heading_color(paragraph, hex_color: str) -> None:
    r, g, b = _hex_to_rgb(hex_color)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(r, g, b)


def _add_colored_heading(doc: Document, text: str, level: int, hex_color: str) -> None:
    p = doc.add_heading(text, level=level)
    _set_heading_color(p, hex_color)


def _add_risk_row(table, risk: dict, row_idx: int) -> None:
    row = table.add_row()
    severity_colors = {"High": CENTRICA_PURPLE, "Medium": CENTRICA_LAVENDER, "Low": CENTRICA_MINT}
    cells = [risk.get("supplier", ""), risk.get("type", ""), risk.get("description", ""), risk.get("severity", "")]
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(9)
                if i == 3:
                    r, g, b = _hex_to_rgb(severity_colors.get(val, CENTRICA_LAVENDER))
                    run.font.color.rgb = RGBColor(r, g, b)
                    run.font.bold = True


def _generate_weekly_digest(category: str, risks: list, news_items: list, run_date: str = "") -> str:
    """Plain function — generate the Word digest. Used by tests and the @function_tool wrapper."""
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    date_str = run_date or datetime.now().strftime("%Y-%m-%d")
    filename = f"{date_str}_digest.docx"
    filepath = OUTPUTS_DIR / filename

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title block
    title = doc.add_heading("", 0)
    title_run = title.add_run("Market Watcher — Weekly Intelligence Digest")
    title_run.font.name = "Arial"
    title_run.font.bold = True
    title_run.font.size = Pt(20)
    r, g, b = _hex_to_rgb(CENTRICA_NAVY)
    title_run.font.color.rgb = RGBColor(r, g, b)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"{category}  ·  {date_str}  ·  Prepared by Market Watcher Agent")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(11)
    r, g, b = _hex_to_rgb(CENTRICA_LAVENDER)
    sub_run.font.color.rgb = RGBColor(r, g, b)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Executive summary
    _add_colored_heading(doc, "Executive Summary", 1, CENTRICA_NAVY)
    high_risks = [r for r in risks if r.get("severity") == "High"]
    med_risks = [r for r in risks if r.get("severity") == "Medium"]
    summary_text = (
        f"This digest covers the {category} supplier category for the week of {date_str}. "
        f"The Market Watcher agent identified {len(high_risks)} HIGH and {len(med_risks)} MEDIUM priority risk signals "
        f"across {len(set(r.get('supplier','') for r in risks))} suppliers, drawing on {len(news_items)} news sources. "
        f"Immediate attention is required on: {'; '.join(r['description'][:60] + '...' for r in high_risks[:2]) if high_risks else 'no critical items this week'}."
    )
    p = doc.add_paragraph(summary_text)
    p.runs[0].font.name = "Calibri"
    p.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # Risk table
    _add_colored_heading(doc, "Risk Register — This Week", 1, CENTRICA_NAVY)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(["Supplier", "Risk Type", "Description", "Severity"]):
        hdr_cells[i].text = hdr
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9)
                r_c, g_c, b_c = _hex_to_rgb(CENTRICA_NAVY)
                run.font.color.rgb = RGBColor(r_c, g_c, b_c)

    for idx, risk in enumerate(risks):
        _add_risk_row(table, risk, idx)

    doc.add_paragraph()

    # News & intelligence
    _add_colored_heading(doc, "Market Intelligence — News Sources", 1, CENTRICA_NAVY)
    for item in news_items[:8]:
        p_head = doc.add_paragraph(style="List Bullet")
        run_head = p_head.add_run(item.get("headline", ""))
        run_head.font.bold = True
        run_head.font.name = "Calibri"
        run_head.font.size = Pt(10)
        r_c, g_c, b_c = _hex_to_rgb(CENTRICA_NAVY)
        run_head.font.color.rgb = RGBColor(r_c, g_c, b_c)

        p_body = doc.add_paragraph()
        snippet = item.get("summary", item.get("snippet", ""))[:200]
        run_body = p_body.add_run(f"{item.get('source', '')}  ·  {item.get('url', '')}  ·  {snippet}")
        run_body.font.name = "Calibri"
        run_body.font.size = Pt(9)
        run_body.font.italic = True

    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    r_c, g_c, b_c = _hex_to_rgb(CENTRICA_PALE_LAV)
    footer_run = footer_p.add_run(
        "CONFIDENTIAL — Centrica Procurement  ·  Generated by Market Watcher Agent  ·  For internal use only"
    )
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(r_c, g_c, b_c)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(filepath)
    return str(filepath)


@function_tool(strict_mode=False)
def generate_weekly_digest(category: str, risks: list, news_items: list, run_date: str = "") -> str:
    """Generate a Centrica-branded Word document weekly intelligence digest.
    Returns the path to the generated .docx file.
    risks: list of dicts with keys: supplier, type, description, severity.
    news_items: list of dicts with keys: headline, source, url, summary.
    """
    return _generate_weekly_digest(category, risks, news_items, run_date)


def _draft_cert_chase_email(supplier_name: str, cert_type: str, cert_id: str, expiry_date: str, days_to_expiry: int, language: str = "english") -> str:
    """Plain function — draft a cert chase email. Used by tests and the @function_tool wrapper."""
    if language.lower() == "spanish":
        subject = f"Aviso: Renovación de Certificación Requerida — {cert_type} ({cert_id})"
        body = f"""Asunto: {subject}

Estimado equipo de {supplier_name},

Nos ponemos en contacto con ustedes en nombre de Centrica Procurement para notificarles que la certificación {cert_type} (referencia: {cert_id}) asociada a su organización vence el {expiry_date} — dentro de {days_to_expiry} días.

De acuerdo con los requisitos de nuestra política de gestión de proveedores, todos los proveedores aprobados deben mantener certificaciones vigentes como condición para continuar en nuestra lista de proveedores preferidos (PSL).

Les solicitamos amablemente que:
1. Inicien el proceso de renovación de la certificación {cert_type} antes del {expiry_date}.
2. Proporcionen el nuevo certificado a nuestro equipo de gestión de proveedores en un plazo máximo de 30 días naturales tras la renovación.
3. Confirmen la recepción de este aviso respondiendo a este correo antes del próximo viernes.

Si tienen alguna duda o necesitan asistencia, no duden en ponerse en contacto con nosotros.

Atentamente,
Equipo de Gestión de Proveedores — Centrica Procurement
[Este correo ha sido generado por el agente Market Watcher — sólo para revisión interna]
"""
    else:
        subject = f"Action Required: Certification Renewal — {cert_type} ({cert_id})"
        urgency = "URGENT — " if days_to_expiry <= 30 else ""
        body = f"""Subject: {urgency}{subject}

Dear {supplier_name} Supplier Relations Team,

I am writing on behalf of Centrica Procurement to notify you that your {cert_type} certification (reference: {cert_id}) is due to expire on {expiry_date} — in {days_to_expiry} {'day' if days_to_expiry == 1 else 'days'}.

As an approved supplier on the Centrica Preferred Supplier List, maintaining current certifications is a contractual requirement. Failure to renew may result in suspension of purchasing activity under your current agreement.

We respectfully request that you:

1. Initiate the {cert_type} renewal process immediately to ensure continuity before {expiry_date}.
2. Forward the renewed certificate to our Supplier Management team within 30 calendar days of issue.
3. Confirm receipt of this notice by return email by end of this week.

If you have already commenced renewal or have an updated certificate available, please send it directly to procurement@centrica.com (mock) quoting reference {cert_id}.

Should you require any support, please do not hesitate to get in touch.

Kind regards,
Centrica Procurement — Supplier Management Team
[This email was drafted by the Market Watcher Agent — for internal review before sending]
"""
    return body


@function_tool
def draft_cert_chase_email(supplier_name: str, cert_type: str, cert_id: str, expiry_date: str, days_to_expiry: int, language: str = "english") -> str:
    """Draft a certification expiry chase email for a supplier.
    language: 'english' or 'spanish' (for demo variety).
    Returns the email body text.
    """
    return _draft_cert_chase_email(supplier_name, cert_type, cert_id, expiry_date, days_to_expiry, language)
